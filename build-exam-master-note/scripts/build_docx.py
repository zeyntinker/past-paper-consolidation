#!/usr/bin/env python3
from __future__ import annotations

import argparse
import sys
from pathlib import Path
from typing import Any, Iterable

from common import load_json
from validate_ledger import validate_ledger


PAGE_WIDTH_DXA = 11906
PAGE_HEIGHT_DXA = 16838
SIDE_MARGIN_DXA = 567  # 10 mm
TOP_BOTTOM_MARGIN_DXA = 1134  # 20 mm
TABLE_INDENT_DXA = 0
TABLE_WIDTH_DXA = 10772
CATEGORY_WIDTH_DXA = 1293  # 12%
PROBLEM_WIDTH_DXA = 4632  # 43%
EXPLANATION_WIDTH_DXA = 4847  # 45%
CELL_MARGIN_TOP_BOTTOM_DXA = 70
CELL_MARGIN_LEFT_RIGHT_DXA = 85
FONT_SIZE_PT = 7
PROVENANCE_SIZE_PT = FONT_SIZE_PT
BODY_SIZE_PT = FONT_SIZE_PT


def require_docx() -> None:
    try:
        import docx  # noqa: F401
    except ImportError as exc:
        raise RuntimeError("python-docx is required to build the master note") from exc


def set_run_font(run: Any, name: str = "Malgun Gothic", size: float | None = None, **kwargs: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_pr.rFonts.set(qn("w:ascii"), name)
    r_pr.rFonts.set(qn("w:hAnsi"), name)
    r_pr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
        sz_cs = r_pr.find(qn("w:szCs"))
        if sz_cs is None:
            sz_cs = OxmlElement("w:szCs")
            r_pr.append(sz_cs)
        sz_cs.set(qn("w:val"), str(int(size * 2)))
    for key, value in kwargs.items():
        setattr(run.font, key, value)


def set_cell_shading(cell: Any, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_width(element: Any, tag: str, width: int) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    target = element.find(qn(tag))
    if target is None:
        target = OxmlElement(tag)
        element.append(target)
    target.set(qn("w:type"), "dxa")
    target.set(qn("w:w"), str(width))


def set_table_caption(table: Any, caption: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl_pr = table._tbl.tblPr
    node = tbl_pr.find(qn("w:tblCaption"))
    if node is None:
        node = OxmlElement("w:tblCaption")
        tbl_pr.append(node)
    node.set(qn("w:val"), caption)


def set_table_geometry(
    table: Any,
    widths_dxa: list[int],
    indent_dxa: int = TABLE_INDENT_DXA,
) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    total_width = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    _set_width(tbl_pr, "w:tblW", total_width)
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for side, value in (
        ("top", CELL_MARGIN_TOP_BOTTOM_DXA),
        ("bottom", CELL_MARGIN_TOP_BOTTOM_DXA),
        ("start", CELL_MARGIN_LEFT_RIGHT_DXA),
        ("end", CELL_MARGIN_LEFT_RIGHT_DXA),
    ):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table._tbl.tr_lst:
        grid_index = 0
        for tc in row.tc_lst:
            tc_pr = tc.get_or_add_tcPr()
            span_node = tc_pr.find(qn("w:gridSpan"))
            span = int(span_node.get(qn("w:val"), "1")) if span_node is not None else 1
            width = sum(widths_dxa[grid_index : grid_index + span])
            _set_width(tc_pr, "w:tcW", width)
            grid_index += span


def repeat_table_header(row: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def prevent_row_split(row: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def clear_container(container: Any) -> None:
    if hasattr(container, "_tc"):
        container.text = ""


def next_paragraph(container: Any) -> Any:
    if hasattr(container, "_tc") and len(container.paragraphs) == 1 and not container.paragraphs[0].text:
        return container.paragraphs[0]
    return container.add_paragraph()


def style_paragraph(paragraph: Any, *, after: float = 2, keep: bool = True) -> None:
    from docx.shared import Pt

    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.keep_together = keep


def configure_section(section: Any) -> None:
    from docx.shared import Twips

    section.page_width = Twips(PAGE_WIDTH_DXA)
    section.page_height = Twips(PAGE_HEIGHT_DXA)
    section.top_margin = Twips(TOP_BOTTOM_MARGIN_DXA)
    section.bottom_margin = Twips(TOP_BOTTOM_MARGIN_DXA)
    section.left_margin = Twips(SIDE_MARGIN_DXA)
    section.right_margin = Twips(SIDE_MARGIN_DXA)
    section.header_distance = Twips(567)
    section.footer_distance = Twips(567)


def clear_story(story: Any) -> None:
    for paragraph in story.paragraphs:
        for child in list(paragraph._p):
            paragraph._p.remove(child)


def set_running_header(section: Any, text: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    section.header.is_linked_to_previous = False
    clear_story(section.header)
    paragraph = section.header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_paragraph(paragraph, after=0)
    run = paragraph.add_run(text)
    set_run_font(run, size=FONT_SIZE_PT, bold=True)
    section.footer.is_linked_to_previous = False
    clear_story(section.footer)


def configure_styles(document: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    for section in document.sections:
        configure_section(section)
        clear_story(section.footer)
        clear_story(section.header)

    tokens = {
        "Normal": (FONT_SIZE_PT, "000000", 0, 2, 1.0),
        "Heading 1": (FONT_SIZE_PT, "000000", 6, 3, 1.0),
        "Heading 2": (FONT_SIZE_PT, "000000", 4, 2, 1.0),
        "Heading 3": (FONT_SIZE_PT, "000000", 3, 2, 1.0),
    }
    for name, (size, color, before, after, spacing) in tokens.items():
        style = document.styles[name]
        style.font.name = "Malgun Gothic"
        r_pr = style._element.get_or_add_rPr()
        r_pr.rFonts.set(qn("w:ascii"), "Malgun Gothic")
        r_pr.rFonts.set(qn("w:hAnsi"), "Malgun Gothic")
        r_pr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        sz_cs = r_pr.find(qn("w:szCs"))
        if sz_cs is None:
            sz_cs = OxmlElement("w:szCs")
            r_pr.append(sz_cs)
        sz_cs.set(qn("w:val"), str(int(size * 2)))
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = spacing


def add_section(document: Any, header_text: str) -> Any:
    from docx.enum.section import WD_SECTION

    section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section)
    set_running_header(section, header_text)
    return section


def create_template(path: Path) -> None:
    require_docx()
    from docx import Document

    document = Document()
    configure_styles(document)
    document.save(path)


def provenance_label(
    provenance: list[dict[str, Any]],
    questions: dict[str, dict[str, Any]],
    sources: dict[str, dict[str, Any]],
    representatives: dict[str, dict[str, Any]] | None = None,
) -> str:
    representatives = representatives or {}
    labels: list[str] = []
    for source in provenance:
        kind = source.get("kind")
        if kind == "original":
            question = questions.get(source.get("source_question_id"), {})
            artifact = sources.get(question.get("source_artifact_id"), {})
            labels.append(
                f"[원문 · {artifact.get('file_name', '?')} · {question.get('source_page', '?')}쪽]"
            )
        elif kind == "ai_reconstruction_from_questions":
            linked = [questions.get(question_id, {}) for question_id in source.get("source_question_ids", [])]
            years = [str(question.get("year")) + "년" for question in linked if question.get("year")]
            files = [
                sources.get(question.get("source_artifact_id"), {}).get("file_name")
                for question in linked
                if question.get("source_artifact_id")
            ]
            basis = "·".join(dict.fromkeys(years or [item for item in files if item])) or "기출"
            labels.append(f"[기출 통합 재구성 · {basis} 기출 참고]")
        elif kind == "lecture_note_supplement":
            artifact = sources.get(source.get("source_artifact_id"), {})
            labels.append(f"[족보 참고 · {artifact.get('file_name', '?')} · {source.get('source_page', '?')}쪽]")
        elif kind == "master_note_supplement":
            representative = representatives.get(source.get("representative_type_id"), {})
            labels.append(
                f"[단권화 보충 · {representative.get('lecture_unit', '?')} · {representative.get('title', '?')}]"
            )
        elif kind == "external_ai_supplement":
            labels.append(f"[AI 외부 보충 · {source.get('citation', '?')} · {source.get('locator', '?')}]")
        else:
            labels.append("[출처 확인 필요]")
    return " ".join(dict.fromkeys(labels))


def add_plain_label(paragraph: Any, text: str) -> None:
    if not text:
        return
    run = paragraph.add_run(f"\n{text}")
    set_run_font(run, size=PROVENANCE_SIZE_PT, italic=True)


def add_component(
    container: Any,
    component: dict[str, Any],
    questions: dict[str, dict[str, Any]],
    sources: dict[str, dict[str, Any]],
    representatives: dict[str, dict[str, Any]] | None = None,
    prefix: str = "",
) -> None:
    paragraph = next_paragraph(container)
    style_paragraph(paragraph)
    if prefix:
        lead = paragraph.add_run(prefix)
        set_run_font(lead, size=BODY_SIZE_PT, bold=True)
    run = paragraph.add_run(component.get("text", ""))
    set_run_font(run, size=BODY_SIZE_PT)
    add_plain_label(
        paragraph,
        provenance_label(component.get("provenance", []), questions, sources, representatives),
    )


def add_labeled_text(container: Any, label: str, text: str, source_label: str = "") -> None:
    if text == "":
        return
    paragraph = next_paragraph(container)
    style_paragraph(paragraph)
    lead = paragraph.add_run(f"{label}: ")
    set_run_font(lead, size=FONT_SIZE_PT, bold=True)
    value = paragraph.add_run(text)
    set_run_font(value, size=BODY_SIZE_PT)
    add_plain_label(paragraph, source_label)


def image_source_label(image: dict[str, Any], sources: dict[str, dict[str, Any]]) -> str:
    artifact = sources.get(image.get("source_artifact_id"), {})
    provenance = image.get("provenance", [])
    if any(item.get("kind") == "original" for item in provenance):
        return f"[원문 · {artifact.get('file_name', '?')} · {image.get('source_page', '?')}쪽]"
    return "[출처 확인 필요]"


def add_image_to_cell(cell: Any, image: dict[str, Any], sources: dict[str, dict[str, Any]]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    path = Path(image.get("path", ""))
    artifact = sources.get(image.get("source_artifact_id"), {})
    paragraph = next_paragraph(cell)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(paragraph, after=3)
    if path.exists():
        run = paragraph.add_run()
        inline_shape = run.add_picture(str(path), width=Inches(6.55))
        alt_text = f"문제 관련 이미지 · {image_source_label(image, sources)}"
        inline_shape._inline.docPr.set("descr", alt_text)
        inline_shape._inline.docPr.set("title", "문제 관련 이미지")
    else:
        run = paragraph.add_run(
            f"[이미지 파일 없음 · {artifact.get('file_name', '?')} · {image.get('source_page', '?')}쪽]"
        )
        set_run_font(run, size=BODY_SIZE_PT, bold=True)
    caption = cell.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(caption, after=2)
    cap_run = caption.add_run(image_source_label(image, sources))
    set_run_font(cap_run, size=PROVENANCE_SIZE_PT, italic=True)


def add_table_image_rows(
    table: Any,
    image_ids: Iterable[str],
    images: dict[str, dict[str, Any]],
    sources: dict[str, dict[str, Any]],
) -> None:
    seen: set[str] = set()
    for image_id in image_ids:
        if image_id in seen or image_id not in images:
            continue
        seen.add(image_id)
        cells = table.add_row().cells
        merged = cells[0].merge(cells[1]).merge(cells[2])
        clear_container(merged)
        add_image_to_cell(merged, images[image_id], sources)


def make_problem_table(document: Any, caption: str) -> Any:
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_caption(table, caption)
    headers = ["구분 / 연도", "문제", "해설"]
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = 1
        style_paragraph(paragraph, after=0)
        run = paragraph.add_run(header)
        set_run_font(run, size=FONT_SIZE_PT, bold=True)
        set_cell_shading(cell, "D9D9D9")
    repeat_table_header(table.rows[0])
    set_table_geometry(table, [CATEGORY_WIDTH_DXA, PROBLEM_WIDTH_DXA, EXPLANATION_WIDTH_DXA])
    return table


def add_category_cell(cell: Any, text: str, *, representative: bool = False) -> None:
    clear_container(cell)
    paragraph = next_paragraph(cell)
    style_paragraph(paragraph, after=0)
    run = paragraph.add_run(text)
    set_run_font(run, size=FONT_SIZE_PT, bold=representative)
    if representative:
        set_cell_shading(cell, "EFEFEF")


def add_occurrence_row(
    table: Any,
    question: dict[str, Any],
    sources: dict[str, dict[str, Any]],
    images: dict[str, dict[str, Any]],
    *,
    category: str,
    supplements: list[dict[str, Any]] | None = None,
    conflict_status: str | None = None,
    representatives: dict[str, dict[str, Any]] | None = None,
) -> None:
    artifact = sources.get(question.get("source_artifact_id"), {})
    cells = table.add_row().cells
    add_category_cell(cells[0], category)
    clear_container(cells[1])
    clear_container(cells[2])
    source_label = exact_source_label(question, artifact)
    add_labeled_text(cells[1], "문제", question.get("original_problem", ""), source_label)
    for index, choice in enumerate(question.get("original_choices", []), start=1):
        add_labeled_text(cells[1], f"선지 {index}", choice, source_label)
    add_labeled_text(cells[2], "원문 답", question.get("original_answer", ""), source_label)
    add_labeled_text(cells[2], "원문 해설", question.get("original_explanation", ""), source_label)
    if supplements:
        label = next_paragraph(cells[2])
        style_paragraph(label, after=2)
        run = label.add_run("검토·보충")
        set_run_font(run, size=FONT_SIZE_PT, bold=True)
        for component in supplements:
            add_component(cells[2], component, {}, sources, representatives)
    if conflict_status not in {None, "none", "resolved"}:
        warning = next_paragraph(cells[2])
        style_paragraph(warning, after=2)
        run = warning.add_run(f"[검토 필요] 원문 답과 보충 판정 불일치 · {conflict_status}")
        set_run_font(run, size=FONT_SIZE_PT, bold=True)
    add_table_image_rows(table, question.get("image_ids", []), images, sources)


def add_representative_table(
    document: Any,
    representative: dict[str, Any],
    questions: dict[str, dict[str, Any]],
    sources: dict[str, dict[str, Any]],
    images: dict[str, dict[str, Any]],
    representatives: dict[str, dict[str, Any]],
) -> None:
    table = make_problem_table(document, f"representative:{representative.get('id', '?')}")
    category, left, right = table.add_row().cells
    add_category_cell(category, f"[대표유형 - {representative.get('question_type', '문제')}]", representative=True)
    clear_container(left)
    clear_container(right)
    set_cell_shading(left, "EFEFEF")
    set_cell_shading(right, "EFEFEF")

    for component in representative.get("problem_components", []):
        add_component(left, component, questions, sources, representatives)
    for index, component in enumerate(representative.get("choice_components", []), start=1):
        add_component(left, component, questions, sources, representatives, prefix=f"{index}. ")

    q_type = representative.get("question_type", "")
    heading = "정답 및 선지별 완전 해설" if q_type in {"객관식", "objective", "multiple_choice"} else "완전 답안 및 해설"
    head = next_paragraph(right)
    style_paragraph(head, after=5)
    head_run = head.add_run(heading)
    set_run_font(head_run, size=FONT_SIZE_PT, bold=True)

    for component in representative.get("answer_components", []):
        add_component(right, component, questions, sources, representatives, prefix="정답: ")
    if q_type in {"객관식", "objective", "multiple_choice"}:
        for index, component in enumerate(representative.get("choice_components", []), start=1):
            verdict = component.get("verdict", "?")
            add_component(right, component, questions, sources, representatives, prefix=f"선지 {index} · {verdict}: ")
            for rationale in component.get("rationale_components", []):
                add_component(right, rationale, questions, sources, representatives, prefix="근거: ")
    for component in representative.get("explanation_components", []):
        add_component(right, component, questions, sources, representatives)

    linked = [questions[qid] for qid in representative.get("question_ids", []) if qid in questions]
    linked.sort(key=lambda item: (item.get("year") or 9999, item.get("source_order") or 9999))
    representative_image_ids = list(
        dict.fromkeys(
            [image_id for image_id in representative.get("image_ids", []) if image_id in images]
            + [
                image_id
                for question in linked
                for image_id in question.get("image_ids", [])
                if image_id in images
            ]
        )
    )
    add_table_image_rows(table, representative_image_ids, images, sources)
    for question in linked:
        artifact = sources.get(question.get("source_artifact_id"), {})
        category_text = " · ".join(
            filter(None, [str(question.get("year", "?")) + "년", artifact.get("file_name")])
        )
        add_occurrence_row(
            table,
            question,
            sources,
            images,
            category=category_text,
            representatives=representatives,
        )
    char_count = sum(len(str(component.get("text", ""))) for component in representative.get("explanation_components", []))
    if char_count < 1000:
        prevent_row_split(table.rows[1])
    set_table_geometry(table, [CATEGORY_WIDTH_DXA, PROBLEM_WIDTH_DXA, EXPLANATION_WIDTH_DXA])


def exact_source_label(question: dict[str, Any], artifact: dict[str, Any]) -> str:
    return f"[원문 · {artifact.get('file_name', '?')} · {question.get('source_page', '?')}쪽]"


def add_source_occurrence(
    document: Any,
    question: dict[str, Any],
    sources: dict[str, dict[str, Any]],
    images: dict[str, dict[str, Any]],
    *,
    location: str,
    supplements: list[dict[str, Any]] | None = None,
    conflict_status: str | None = None,
    representatives: dict[str, dict[str, Any]] | None = None,
) -> None:
    from docx.shared import Pt

    artifact = sources.get(question.get("source_artifact_id"), {})
    table = make_problem_table(document, f"occurrence:{location}:{question.get('id', '?')}")
    sequence = location.rsplit(":", 1)[-1] if location.startswith("part2:") else ""
    category = " · ".join(
        filter(
            None,
            [
                f"작년 · {sequence}번" if sequence else None,
                str(question.get("year", "?")) + "년",
                artifact.get("file_name"),
            ],
        )
    )
    add_occurrence_row(
        table,
        question,
        sources,
        images,
        category=category,
        supplements=supplements,
        conflict_status=conflict_status,
        representatives=representatives,
    )
    source_length = sum(
        len(str(value))
        for value in (
            question.get("original_problem", ""),
            question.get("original_answer", ""),
            question.get("original_explanation", ""),
            *question.get("original_choices", []),
        )
    )
    if source_length < 1000 and len(table.rows) > 1:
        prevent_row_split(table.rows[1])
    set_table_geometry(table, [CATEGORY_WIDTH_DXA, PROBLEM_WIDTH_DXA, EXPLANATION_WIDTH_DXA])
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_title_page(document: Any, ledger: dict[str, Any], status: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(12)
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker_run = kicker.add_run("VERIFIED EXAM REFERENCE")
    set_run_font(kicker_run, size=FONT_SIZE_PT, bold=True)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(ledger.get("course_title", "시험 단권화 노트"))
    set_run_font(title_run, size=FONT_SIZE_PT, bold=True)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(f"강의 순서 단권화 · 작년 기출 실전 순서 · 완전성 감사표\n상태: {status}")
    set_run_font(subtitle_run, size=FONT_SIZE_PT)


def humanize_audit_text(
    value: object,
    sources: dict[str, dict[str, Any]],
    questions: dict[str, dict[str, Any]],
    images: dict[str, dict[str, Any]],
    representatives: dict[str, dict[str, Any]],
) -> str:
    text = str(value)
    replacements: dict[str, str] = {}
    for source_id, source in sources.items():
        replacements[source_id] = source.get("file_name", "원본 파일")
    for question_id, question in questions.items():
        artifact = sources.get(question.get("source_artifact_id"), {})
        replacements[question_id] = (
            f"{question.get('year', '?')}년 {artifact.get('file_name', '?')} {question.get('source_page', '?')}쪽"
        )
    for image_id, item in images.items():
        artifact = sources.get(item.get("source_artifact_id"), {})
        replacements[image_id] = f"{artifact.get('file_name', '?')} {item.get('source_page', '?')}쪽 이미지"
    for rep_id, representative in representatives.items():
        replacements[rep_id] = f"{representative.get('lecture_unit', '')} {representative.get('title', '')}".strip()
    for internal_id in sorted(replacements, key=len, reverse=True):
        text = text.replace(internal_id, replacements[internal_id])
    return text


def add_compact_audit(
    document: Any,
    ledger: dict[str, Any],
    validation: dict[str, Any],
    sources: dict[str, dict[str, Any]],
    questions: dict[str, dict[str, Any]],
    images: dict[str, dict[str, Any]],
    representatives: dict[str, dict[str, Any]],
) -> None:
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    set_table_caption(table, "audit:coverage")
    headers = ["대표문제", "문제 완전성", "해설 완전성", "미해결", "상태"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
        set_cell_shading(table.rows[0].cells[index], "D9D9D9")
    repeat_table_header(table.rows[0])
    for summary in validation.get("representative_coverage", []):
        cells = table.add_row().cells
        values = [
            " · ".join(filter(None, [summary.get("lecture_unit"), summary.get("title")])),
            f"{summary.get('problem_mapped', 0)}/{summary.get('problem_required', 0)}",
            f"{summary.get('explanation_mapped', 0)}/{summary.get('explanation_required', 0)}",
            summary.get("unresolved", 0),
            "완료" if summary.get("status") == "complete" else "검수 필요",
        ]
        for index, value in enumerate(values):
            cells[index].text = str(value)
            for paragraph in cells[index].paragraphs:
                style_paragraph(paragraph)
                for run in paragraph.runs:
                    set_run_font(run, size=BODY_SIZE_PT)
    set_table_geometry(table, [3300, 1740, 1740, 1200, 1860])

    review_summaries = [
        summary
        for summary in validation.get("representative_coverage", [])
        if summary.get("status") != "complete"
    ]
    findings = [
        finding
        for finding in ledger.get("audit_findings", [])
        if finding.get("severity") == "blocking" and finding.get("status") != "resolved"
    ]
    if not findings and not review_summaries:
        return
    heading = document.add_heading("검수 필요 사항", level=2)
    heading.paragraph_format.keep_with_next = True
    for summary in review_summaries:
        paragraph = document.add_paragraph()
        style_paragraph(paragraph)
        detail = (
            f"{summary.get('lecture_unit', '')} · {summary.get('title', '')}: "
            f"문제 의미 반영 {summary.get('problem_mapped', 0)}/{summary.get('problem_required', 0)}, "
            f"해설 의미 반영 {summary.get('explanation_mapped', 0)}/{summary.get('explanation_required', 0)}, "
            f"미해결 {summary.get('unresolved', 0)}건"
        )
        run = paragraph.add_run(detail)
        set_run_font(run, size=BODY_SIZE_PT)
    for finding in findings:
        locations = ", ".join(
            humanize_audit_text(item, sources, questions, images, representatives)
            for item in finding.get("source_locations", [])
        )
        message = humanize_audit_text(
            finding.get("message", "확인이 필요합니다."), sources, questions, images, representatives
        )
        paragraph = document.add_paragraph()
        style_paragraph(paragraph)
        run = paragraph.add_run(message + (f" · {locations}" if locations else ""))
        set_run_font(run, size=BODY_SIZE_PT)


def build_document(ledger: dict[str, Any], template: Path | None, status: str) -> Any:
    require_docx()
    from docx import Document

    document = Document(str(template)) if template and template.exists() else Document()
    configure_styles(document)
    add_title_page(document, ledger, status)

    sources = {item["id"]: item for item in ledger.get("source_artifacts", [])}
    questions = {item["id"]: item for item in ledger.get("question_occurrences", [])}
    images = {item["id"]: item for item in ledger.get("images", [])}
    representatives = ledger.get("representative_types", [])
    representative_map = {item["id"]: item for item in representatives}
    validation = validate_ledger(ledger)

    first_header = representatives[0].get("lecture_unit", "강의 순서 단권화") if representatives else "강의 순서 단권화"
    add_section(document, first_header)
    document.add_heading("제1부. 강의 순서 단권화", level=1)
    current_unit = first_header
    for rep_index, representative in enumerate(representatives):
        lecture_unit = representative.get("lecture_unit", "강의 순서 단권화")
        if rep_index and lecture_unit != current_unit:
            add_section(document, lecture_unit)
            current_unit = lecture_unit
        document.add_heading(
            f"{lecture_unit} · {representative.get('title', '')}",
            level=2,
        ).paragraph_format.keep_with_next = True
        document.add_heading("대표 문제와 완전 해설", level=3).paragraph_format.keep_with_next = True
        add_representative_table(document, representative, questions, sources, images, representative_map)

    add_section(document, "작년 기출 실전 순서")
    document.add_heading("제2부. 작년 기출 실전 순서", level=1)
    for entry in ledger.get("prior_year_sequence", []):
        for question_id in entry.get("question_ids", []):
            if question_id in questions:
                add_source_occurrence(
                    document,
                    questions[question_id],
                    sources,
                    images,
                    location=f"part2:{entry.get('sequence')}",
                    supplements=entry.get("supplement_components", []),
                    conflict_status=entry.get("conflict_status"),
                    representatives=representative_map,
                )

    add_section(document, "완전성 감사표")
    document.add_heading("제3부. 완전성 감사표", level=1)
    add_compact_audit(
        document,
        ledger,
        validation,
        sources,
        questions,
        images,
        representative_map,
    )
    return document


def main() -> int:
    parser = argparse.ArgumentParser(description="Build a three-part, provenance-aware exam DOCX")
    parser.add_argument("ledger", nargs="?")
    parser.add_argument("--template")
    parser.add_argument("--out")
    parser.add_argument("--review-draft", action="store_true")
    parser.add_argument("--create-template")
    args = parser.parse_args()
    if args.create_template:
        create_template(Path(args.create_template))
        return 0
    if not args.ledger or not args.out:
        parser.error("ledger and --out are required unless --create-template is used")
    ledger = load_json(args.ledger)
    validation = validate_ledger(ledger)
    if not validation["valid"] and not args.review_draft:
        print("ledger validation failed; use --review-draft only for an explicitly labeled draft")
        return 2
    status = "완성 후보" if validation["valid"] else "검수 필요 초안"
    document = build_document(ledger, Path(args.template) if args.template else None, status)
    output = Path(args.out)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    print(output)
    return 0


if __name__ == "__main__":
    sys.exit(main())
